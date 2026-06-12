import io
import docx
import logging
from pypdf import PdfReader

logger = logging.getLogger(__name__)

def parse_pdf(file_bytes: bytes) -> str:
    """Parses text content from PDF file bytes using pypdf."""
    text_content = []
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        raise ValueError(f"Failed to parse PDF document: {str(e)}")
    
    return "\n".join(text_content)

def parse_docx(file_bytes: bytes) -> str:
    """Parses text content from DOCX file bytes."""
    try:
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract from tables as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
                        
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")

def count_tokens(text: str, model_name: str = "gpt-4") -> int:
    """Counts tokens using tiktoken, falling back to approximation if offline or not installed."""
    try:
        import tiktoken
        # Llama models aren't directly supported by tiktoken, but gpt-4/cl100k_base is a very good proxy
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))
    except Exception as e:
        logger.warning(f"Failed to load/use tiktoken (possibly not installed), falling back to estimation: {e}")
        # Standard fallback approximation: 1 token ~ 4 characters
        return len(text) // 4
